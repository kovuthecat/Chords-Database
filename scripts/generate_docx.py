"""
Génère un fichier .docx à partir d'un JSON song validé.
Usage : python scripts/generate_docx.py data/song_<slug>.json
"""

import io
import json
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config import OUTPUT_DIR, PDF_EXPORT_DIR
from memo import build_memo_lines, apply_substitutions_to_song, _extract_performance_lines


def _ensure_utf8():
    if hasattr(sys.stdout, "buffer") and getattr(sys.stdout, "encoding", "").lower() != "utf-8":
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")


# ---------------------------------------------------------------------------
# Polices et tailles
# ---------------------------------------------------------------------------

MONO_FONT   = "Consolas"       # lignes accords + paroles (monospace = alignement garanti)
BODY_FONT   = "Calibri"        # titres, en-têtes, méta

CHORD_SIZE  = Pt(13)           # accords : légèrement plus grand pour visibilité
LYRIC_SIZE  = Pt(12)           # paroles
SECTION_SIZE= Pt(13)           # [ SECTION ]
TITLE_SIZE  = Pt(22)           # titre chanson
ARTIST_SIZE = Pt(14)           # artiste
META_SIZE   = Pt(10)           # capo, tonalité, etc.

CHORD_COLOR   = RGBColor(0x1F, 0x4E, 0x79)   # bleu marine bien lisible
SECTION_COLOR = RGBColor(0x20, 0x60, 0x20)  # vert foncé pour les en-têtes de section
MEMO_GRAY     = RGBColor(0x55, 0x55, 0x55)  # gris pour répétitions / rhythm hint


# ---------------------------------------------------------------------------
# Chargement et validation
# ---------------------------------------------------------------------------

def load_song(json_path: str) -> dict:
    path = Path(json_path)
    if not path.exists():
        print(f"Erreur : fichier introuvable : {json_path}")
        sys.exit(1)
    with open(path, encoding="utf-8") as f:
        return json.load(f)


def check_validated(song: dict):
    status = song.get("validation", {}).get("status", "pending")
    if status != "user_validated":
        print(f"Erreur : le fichier n'est pas encore validé (status={status}).")
        print("Lance d'abord : python scripts/display_validation.py <fichier>")
        sys.exit(1)


# ---------------------------------------------------------------------------
# Helpers paragraphe
# ---------------------------------------------------------------------------

def _p_keep_with_next(p):
    """Empêche un saut de page entre ce paragraphe et le suivant."""
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.page_break_before = False


def _p_no_space(p):
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)


def _p_small_space(p):
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(0)


# ---------------------------------------------------------------------------
# Blocs de contenu
# ---------------------------------------------------------------------------

def add_title_block(doc: Document, song: dict):
    meta = song["meta"]

    # Titre
    p = doc.add_paragraph()
    run = p.add_run(meta["title"].upper())
    run.bold = True
    run.font.name = BODY_FONT
    run.font.size = TITLE_SIZE
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)

    # Artiste
    p = doc.add_paragraph()
    run = p.add_run(meta["artist"])
    run.font.name = BODY_FONT
    run.font.size = ARTIST_SIZE
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)

    # Album (optionnel)
    if meta.get("album"):
        p = doc.add_paragraph()
        run = p.add_run(meta["album"])
        run.font.name = BODY_FONT
        run.font.size = META_SIZE
        run.italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)
    else:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)

    # Ligne méta : tonalité, capo, tempo
    meta_parts = []
    if meta.get("key"):
        mode = "majeur" if meta.get("key_mode") == "major" else "mineur"
        meta_parts.append(f"Tonalité : {meta['key']} {mode}")
    if meta.get("capo", 0):
        meta_parts.append(f"Capo : {meta['capo']}")
    if meta.get("tempo"):
        meta_parts.append(f"Tempo : ~{meta['tempo']} bpm")
    if meta.get("tuning") and meta["tuning"] != "standard":
        meta_parts.append(f"Accordage : {meta['tuning']}")

    if meta_parts:
        p = doc.add_paragraph("   |   ".join(meta_parts))
        run = p.runs[0]
        run.font.name = BODY_FONT
        run.font.size = META_SIZE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)

    # Accords utilisés
    chords = [c for c in song.get("chords_used", []) if not c.startswith("_")]
    if chords:
        p = doc.add_paragraph("Accords : " + "  ".join(chords))
        run = p.runs[0]
        run.font.name = BODY_FONT
        run.font.size = META_SIZE
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(16)


def add_section_header(doc: Document, label: str):
    p = doc.add_paragraph()
    run = p.add_run(f"[ {label.upper()} ]")
    run.bold = True
    run.font.name = BODY_FONT
    run.font.size = SECTION_SIZE
    run.font.color.rgb = SECTION_COLOR
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    # Évite qu'un en-tête de section se retrouve seul en bas de page
    p.paragraph_format.widow_control = True


def add_chord_grid(doc: Document, chord_grid: str, repeats: int = 1):
    """Ajoute une grille d'accords instrumentale. Gère les grilles multi-lignes (\n)."""
    lines = chord_grid.split("\n")
    for i, grid_line in enumerate(lines):
        grid_line = grid_line.strip()
        if not grid_line:
            continue
        # Indicateur ×N seulement sur la dernière ligne
        suffix = f"   ×{repeats}" if (repeats > 1 and i == len(lines) - 1) else ""
        p = doc.add_paragraph(grid_line + suffix)
        run = p.runs[0]
        run.font.name = MONO_FONT
        run.font.size = CHORD_SIZE
        run.bold = True
        run.font.color.rgb = CHORD_COLOR
        _p_keep_with_next(p)
        _p_no_space(p)
        p.paragraph_format.space_after = Pt(2)


MAX_LINE_LEN = 80  # caractères max pour une ligne accords/paroles en Consolas 12pt sur A4

def build_chord_line(chords: list, lyrics: str) -> str:
    """Place les accords à leurs positions en caractères au-dessus des paroles."""
    if not chords:
        return ""
    max_pos = max(c["position"] for c in chords)
    line_len = max(len(lyrics) + 4, max_pos + 8)
    result = [" "] * line_len
    for entry in sorted(chords, key=lambda x: x["position"]):
        pos = entry["position"]
        for i, ch in enumerate(entry["chord"]):
            if pos + i < len(result):
                result[pos + i] = ch
    return "".join(result).rstrip()


def add_chord_lyric_lines(doc: Document, lines: list):
    for line in lines:
        chords = line.get("chords", [])
        lyrics = line.get("lyrics", "")

        # Ligne d'accords
        if chords:
            chord_text = build_chord_line(chords, lyrics)
            p = doc.add_paragraph()
            run = p.add_run(chord_text)
            run.font.name = MONO_FONT
            run.font.size = CHORD_SIZE
            run.bold = True
            run.font.color.rgb = CHORD_COLOR
            _p_no_space(p)
            _p_keep_with_next(p)   # colle au paragraphe suivant (paroles)

        # Ligne de paroles
        p = doc.add_paragraph()
        run = p.add_run(lyrics or "")
        run.font.name = MONO_FONT
        run.font.size = LYRIC_SIZE
        _p_small_space(p)


# ---------------------------------------------------------------------------
# Rendu d'une section
# ---------------------------------------------------------------------------

def render_section(doc: Document, section: dict):
    label = section.get("label", section.get("type", "Section"))
    add_section_header(doc, label)

    repeats = section.get("repeats", 1)
    chord_grid = section.get("chord_grid")
    lines = section.get("lines", [])
    is_instr = section.get("is_instrumental", False)

    if is_instr:
        perf_explicit = section.get("performance_progression")
        if perf_explicit and any(e.get("chords") for e in perf_explicit):
            # performance_progression prioritaire : affiche les groupements avec ×N par ligne
            for entry in perf_explicit:
                chords = str(entry.get("chords") or "")
                if not chords:
                    continue
                suffix = f"  ×{entry['repeat']}" if entry.get("repeat", 1) > 1 else ""
                p = doc.add_paragraph(chords + suffix)
                run = p.runs[0]
                run.font.name = MONO_FONT
                run.font.size = Pt(12)
                run.bold = True
                run.font.color.rgb = CHORD_COLOR
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(3)
        elif chord_grid:
            add_chord_grid(doc, chord_grid, repeats)
        else:
            perf = _extract_performance_lines(section)
            has_chords = any(pl.get("chords") for pl in perf)
            if has_chords:
                for pl in perf:
                    if not pl.get("chords"):
                        continue
                    suffix = f"  ×{pl['repeat']}" if pl.get("repeat", 1) > 1 else ""
                    p = doc.add_paragraph(pl["chords"] + suffix)
                    run = p.runs[0]
                    run.font.name = MONO_FONT
                    run.font.size = Pt(12)
                    run.bold = True
                    run.font.color.rgb = CHORD_COLOR
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(3)
            else:
                p = doc.add_paragraph("(instrumental)")
                run = p.runs[0]
                run.font.name = BODY_FONT
                run.font.size = META_SIZE
                run.italic = True
    else:
        if lines:
            add_chord_lyric_lines(doc, lines)
            if repeats > 1:
                p = doc.add_paragraph(f"(× {repeats})")
                run = p.runs[0]
                run.font.name = BODY_FONT
                run.font.size = META_SIZE
                run.italic = True
        elif chord_grid:
            add_chord_grid(doc, chord_grid, repeats)

    # Espace après la section
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)


# ---------------------------------------------------------------------------
# Fiche mémo structure guitare
# ---------------------------------------------------------------------------

def add_memo_page_docx(doc: Document, song: dict, simplified_song: dict = None,
                       page_break: bool = True):
    """Ajoute la fiche mémo structure (saut de page optionnel pour rendu partiel).
    Si simplified_song est fourni (mode='both'), ajoute un bloc VERSION SIMPLIFIÉE."""
    if page_break:
        doc.add_page_break()

    meta = song["meta"]
    memo_lines = build_memo_lines(song)
    if not memo_lines:
        return

    # Titre
    p = doc.add_paragraph()
    run = p.add_run(f"{meta.get('title', '').upper()}  —  {meta.get('artist', '')}")
    run.bold = True
    run.font.name = BODY_FONT
    run.font.size = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)

    # Sous-titre
    p = doc.add_paragraph()
    run = p.add_run("FICHE MÉMO STRUCTURE")
    run.bold = True
    run.font.name = BODY_FONT
    run.font.size = Pt(10)
    run.font.color.rgb = SECTION_COLOR
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)

    # Capo si présent
    capo = meta.get("capo", 0)
    if capo:
        p = doc.add_paragraph()
        run = p.add_run(f"Capo {capo}")
        run.font.name = BODY_FONT
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)

    # Séparateur haut
    p = doc.add_paragraph("─" * 52)
    p.runs[0].font.name = MONO_FONT
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
    _p_no_space(p)
    p.paragraph_format.space_after = Pt(6)

    # Alignement : largeur max des labels
    max_label = max(len(m["label"]) for m in memo_lines)
    indent = " " * (max_label + 3)  # espace sous le label "Label : "

    # Conducteur guitare — une ou plusieurs lignes par section
    for m in memo_lines:
        lines_data = m.get("lines") or [{"chords": m["progression"] or "—", "repeat": 1}]

        for idx, ln in enumerate(lines_data):
            p = doc.add_paragraph()

            if idx == 0:
                run = p.add_run(m["label"].ljust(max_label) + " : ")
                run.font.name = MONO_FONT
                run.font.size = Pt(11)
                run.bold = True
            else:
                run = p.add_run(indent)
                run.font.name = MONO_FONT
                run.font.size = Pt(11)

            # Accords de la ligne
            run = p.add_run(ln.get("chords") or "—")
            run.font.name = MONO_FONT
            run.font.size = Pt(11)
            run.bold = True
            run.font.color.rgb = CHORD_COLOR

            # Repeat de ligne (ex: (x3))
            ln_repeat = ln.get("repeat", 1)
            if ln_repeat > 1:
                run = p.add_run(f"  (x{ln_repeat})")
                run.font.name = MONO_FONT
                run.font.size = Pt(10)
                run.italic = True
                run.font.color.rgb = MEMO_GRAY

            # Section-level repeat sur la dernière ligne
            if idx == len(lines_data) - 1 and m["repeat"]:
                run = p.add_run(f"  {m['repeat']}")
                run.font.name = MONO_FONT
                run.font.size = Pt(10)
                run.italic = True
                run.font.color.rgb = MEMO_GRAY

            # Rhythm hint sur la première ligne
            if idx == 0 and m["rhythm_hint"]:
                run = p.add_run(f"  — {m['rhythm_hint']}")
                run.font.name = MONO_FONT
                run.font.size = Pt(10)
                run.italic = True
                run.font.color.rgb = MEMO_GRAY

            _p_no_space(p)
            p.paragraph_format.space_after = Pt(2 if idx < len(lines_data) - 1 else 4)

        # Mini-tab : Consolas 9pt, max 2 lignes
        for tab_line in (m["mini_tab"] or [])[:2]:
            p_tab = doc.add_paragraph()
            run = p_tab.add_run("    " + tab_line)
            run.font.name = MONO_FONT
            run.font.size = Pt(9)
            _p_no_space(p_tab)
            p_tab.paragraph_format.space_after = Pt(2)

    # --- Bloc VERSION SIMPLIFIÉE (mode="both") ---
    if simplified_song:
        orig_lines = build_memo_lines(song)
        simp_lines = build_memo_lines(simplified_song)

        # Séparateur intermédiaire
        p = doc.add_paragraph("─" * 52)
        p.runs[0].font.name = MONO_FONT
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
        _p_no_space(p)
        p.paragraph_format.space_after = Pt(4)

        # Titre du bloc simplifié
        p = doc.add_paragraph()
        run = p.add_run("VERSION SIMPLIFIÉE")
        run.bold = True
        run.font.name = BODY_FONT
        run.font.size = Pt(10)
        run.font.color.rgb = SECTION_COLOR
        p.paragraph_format.space_after = Pt(4)

        max_label_s = max((len(m["label"]) for m in simp_lines), default=max_label)
        indent_s = " " * (max_label_s + 3)

        for i, m in enumerate(simp_lines):
            orig_prog = orig_lines[i]["progression"] if i < len(orig_lines) else ""
            identical = (m["progression"] == orig_prog)
            lines_data = m.get("lines") or [{"chords": m["progression"] or "—", "repeat": 1}]

            for idx, ln in enumerate(lines_data):
                p = doc.add_paragraph()

                if idx == 0:
                    run = p.add_run(m["label"].ljust(max_label_s) + " : ")
                    run.font.name = MONO_FONT
                    run.font.size = Pt(11)
                    run.bold = True
                else:
                    run = p.add_run(indent_s)
                    run.font.name = MONO_FONT
                    run.font.size = Pt(11)

                run = p.add_run(ln.get("chords") or "—")
                run.font.name = MONO_FONT
                run.font.size = Pt(11)
                run.bold = True
                run.font.color.rgb = MEMO_GRAY if identical else CHORD_COLOR

                ln_repeat = ln.get("repeat", 1)
                if ln_repeat > 1 and not identical:
                    run = p.add_run(f"  (x{ln_repeat})")
                    run.font.name = MONO_FONT
                    run.font.size = Pt(10)
                    run.italic = True
                    run.font.color.rgb = MEMO_GRAY

                if idx == len(lines_data) - 1 and m["repeat"] and not identical:
                    run = p.add_run(f"  {m['repeat']}")
                    run.font.name = MONO_FONT
                    run.font.size = Pt(10)
                    run.italic = True
                    run.font.color.rgb = MEMO_GRAY

                _p_no_space(p)
                p.paragraph_format.space_after = Pt(2 if idx < len(lines_data) - 1 else 4)

    # Séparateur bas
    p = doc.add_paragraph("─" * 52)
    p.runs[0].font.name = MONO_FONT
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
    _p_no_space(p)


def _add_memo_page_pdf(story: list, song: dict, simplified_song: dict = None,
                       page_break: bool = True):
    """Ajoute la fiche mémo structure dans le story reportlab (saut de page optionnel).
    Si simplified_song est fourni (mode='both'), ajoute un bloc VERSION SIMPLIFIÉE."""
    from reportlab.platypus import PageBreak, Paragraph, Preformatted, HRFlowable
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib import colors
    from reportlab.lib.units import mm
    from reportlab.lib.enums import TA_CENTER

    c_section = colors.HexColor("#206020")
    c_hr      = colors.HexColor("#BBBBBB")

    def esc(txt):
        return txt.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    memo_lines = build_memo_lines(song)
    if not memo_lines:
        return

    if page_break:
        story.append(PageBreak())

    meta = song["meta"]
    story.append(Paragraph(
        esc(f"{meta.get('title', '').upper()}  —  {meta.get('artist', '')}"),
        ParagraphStyle("mt", fontName="Helvetica-Bold", fontSize=16,
                       textColor=colors.HexColor("#1F2D5A"),
                       alignment=TA_CENTER, spaceAfter=2*mm),
    ))
    story.append(Paragraph(
        "FICHE MÉMO STRUCTURE",
        ParagraphStyle("ms", fontName="Helvetica-Bold", fontSize=10,
                       textColor=c_section, alignment=TA_CENTER, spaceAfter=3*mm),
    ))

    capo = meta.get("capo", 0)
    if capo:
        story.append(Paragraph(
            f"Capo {capo}",
            ParagraphStyle("mc", fontName="Helvetica", fontSize=10,
                           alignment=TA_CENTER, spaceAfter=4*mm),
        ))

    story.append(HRFlowable(width="100%", thickness=0.5, color=c_hr, spaceAfter=4*mm))

    max_label = max(len(m["label"]) for m in memo_lines)
    indent = " " * (max_label + 3)
    line_style = ParagraphStyle("ml", fontName="Courier", fontSize=10,
                                spaceAfter=2*mm, leading=13)
    tab_style  = ParagraphStyle("mt2", fontName="Courier", fontSize=9,
                                spaceAfter=1*mm, leading=11)

    for m in memo_lines:
        lines_data = m.get("lines") or [{"chords": m["progression"] or "—", "repeat": 1}]
        is_last_line = lambda idx: idx == len(lines_data) - 1

        for idx, ln in enumerate(lines_data):
            chords = ln.get("chords") or "—"
            ln_repeat = ln.get("repeat", 1)
            repeat_str = f"  (x{ln_repeat})" if ln_repeat > 1 else ""
            section_repeat = f"  {m['repeat']}" if (is_last_line(idx) and m["repeat"]) else ""
            rhythm_str = f"  — {m['rhythm_hint']}" if (m["rhythm_hint"] and is_last_line(idx)) else ""
            prefix = f"{m['label'].ljust(max_label)} : " if idx == 0 else indent
            line_text = f"{prefix}{chords}{repeat_str}{section_repeat}{rhythm_str}"
            story.append(Preformatted(line_text, line_style))

        for tab_line in (m["mini_tab"] or [])[:2]:
            story.append(Preformatted("    " + tab_line, tab_style))

    # Bloc VERSION SIMPLIFIÉE (mode="both")
    if simplified_song:
        orig_lines = build_memo_lines(song)
        simp_lines = build_memo_lines(simplified_song)

        story.append(HRFlowable(width="100%", thickness=0.5, color=c_hr, spaceAfter=2*mm))
        story.append(Paragraph(
            "VERSION SIMPLIFIÉE",
            ParagraphStyle("sv", fontName="Helvetica-Bold", fontSize=10,
                           textColor=c_section, spaceAfter=2*mm),
        ))

        for i, m in enumerate(simp_lines):
            orig_prog = orig_lines[i]["progression"] if i < len(orig_lines) else ""
            identical = (m["progression"] == orig_prog)
            parts = [m["progression"] or "—"]
            if m["repeat"] and not identical:
                parts.append(m["repeat"])
            line_text = f"{m['label'].ljust(max_label)} : {' '.join(parts)}"
            s = ParagraphStyle(
                f"sl{i}", fontName="Courier",
                fontSize=10, spaceAfter=2*mm, leading=13,
                textColor=colors.HexColor("#555555") if identical else colors.HexColor("#1F4E79"),
            )
            story.append(Preformatted(line_text, s))

    story.append(HRFlowable(width="100%", thickness=0.5, color=c_hr, spaceAfter=0))


# ---------------------------------------------------------------------------
# Génération du document
# ---------------------------------------------------------------------------

def set_margins(doc: Document, top=2.0, bottom=2.0, left=2.0, right=2.0):
    """Définit les marges en cm pour toutes les sections du document."""
    for section in doc.sections:
        section.top_margin    = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin   = Cm(left)
        section.right_margin  = Cm(right)


def _resolve_simplification(song: dict) -> tuple:
    """
    Retourne (render_song, simplified_song_or_None) selon le mode de simplification.
    - "original"   → (song, None)
    - "simplified" → (song avec substitutions, None)
    - "both"       → (song original, song avec substitutions)
    """
    simp = song.get("simplification", {})
    mode = simp.get("mode") or "original"
    substitutions = simp.get("chord_substitutions", {})

    if mode == "simplified" and substitutions:
        return (apply_substitutions_to_song(song, substitutions), None)
    elif mode == "both" and substitutions:
        return (song, apply_substitutions_to_song(song, substitutions))
    return (song, None)


def generate(song: dict, output_path: Path):
    render_song, simplified_song = _resolve_simplification(song)

    doc = Document()
    set_margins(doc, top=2.0, bottom=2.0, left=2.5, right=2.0)

    add_title_block(doc, render_song)

    sections_by_id = {s["id"]: s for s in render_song.get("sections", [])}
    sequence = [
        x for x in render_song.get("structure_sequence", [])
        if isinstance(x, str) and not x.startswith("_comment")
    ]

    if sequence:
        for section_id in sequence:
            section = sections_by_id.get(section_id)
            if section:
                render_section(doc, section)
    else:
        for section in render_song.get("sections", []):
            render_section(doc, section)

    # Avertissements critiques en bas de document
    warnings = [w for w in render_song.get("warnings", []) if w.get("severity") in ("medium", "high")]
    if warnings:
        p = doc.add_paragraph("Notes :")
        p.runs[0].bold = True
        p.runs[0].font.name = BODY_FONT
        for w in warnings:
            p = doc.add_paragraph(f"• {w['message']}")
            p.runs[0].font.name = BODY_FONT
            p.runs[0].font.size = META_SIZE

    # Fiche mémo structure — page obligatoire
    add_memo_page_docx(doc, render_song, simplified_song=simplified_song)

    doc.save(output_path)
    print(f"DOCX généré : {output_path}")


# ---------------------------------------------------------------------------
# Génération PDF (reportlab)
# ---------------------------------------------------------------------------

def generate_pdf(song: dict, output_path: Path, simplified_song: dict = None, parts=None):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm, mm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Preformatted, Spacer, HRFlowable, KeepTogether
    )
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_CENTER

    c_title   = colors.HexColor("#1F2D5A")
    c_chord   = colors.HexColor("#1F4E79")
    c_section = colors.HexColor("#206020")
    c_meta    = colors.HexColor("#555555")
    c_hr      = colors.HexColor("#BBBBBB")

    def st(name, **kw):
        return ParagraphStyle(name, **kw)

    S = {
        "title":   st("t",  fontName="Helvetica-Bold",    fontSize=20, textColor=c_title,   alignment=TA_CENTER, spaceAfter=2*mm),
        "artist":  st("ar", fontName="Helvetica",         fontSize=14, textColor=c_title,   alignment=TA_CENTER, spaceAfter=4*mm),
        "meta":    st("m",  fontName="Helvetica",         fontSize=10, textColor=c_meta,    alignment=TA_CENTER, spaceAfter=2*mm),
        "section": st("s",  fontName="Helvetica-Bold",    fontSize=12, textColor=c_section, spaceBefore=6*mm,    spaceAfter=2*mm),
        # chord et lyric : même taille obligatoire pour que les colonnes s'alignent
        "chord":   st("ch", fontName="Courier-Bold", fontSize=10, textColor=c_chord,      spaceAfter=0,    leading=12),
        "lyric":   st("ly", fontName="Courier",      fontSize=10, textColor=colors.black, spaceAfter=3*mm, leading=12),
        "grid":    st("gr", fontName="Courier-Bold", fontSize=10, textColor=c_chord,      spaceAfter=2*mm, leading=12),
        "instr":   st("in", fontName="Helvetica-Oblique", fontSize=9,  textColor=c_meta,    spaceAfter=2*mm),
        "repeat":  st("rp", fontName="Helvetica-Oblique", fontSize=9,  textColor=c_meta,    spaceAfter=2*mm),
        "warn":    st("w",  fontName="Helvetica",         fontSize=9,  textColor=c_meta,    spaceAfter=1*mm),
    }

    def esc(txt):
        return txt.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    meta   = song["meta"]
    _parts = set(parts) if parts else {"chords", "memo"}
    story  = []

    story.append(Paragraph(esc(meta["title"].upper()), S["title"]))
    story.append(Paragraph(esc(meta["artist"]),         S["artist"]))
    if meta.get("album"):
        story.append(Paragraph(esc(meta["album"]),
                                ParagraphStyle("alb", fontName="Helvetica-Oblique", fontSize=10,
                                               textColor=c_meta, alignment=TA_CENTER, spaceAfter=2*mm)))

    meta_parts = []
    if meta.get("key"):
        mode = "majeur" if meta.get("key_mode") == "major" else "mineur"
        meta_parts.append(f"Tonalité : {meta['key']} {mode}")
    if meta.get("capo"):
        meta_parts.append(f"Capo : {meta['capo']}")
    if meta.get("tempo"):
        meta_parts.append(f"♩ = {meta['tempo']} bpm")
    if meta_parts:
        story.append(Paragraph("   |   ".join(meta_parts), S["meta"]))

    chords = [c for c in song.get("chords_used", []) if not c.startswith("_")]
    if chords:
        story.append(Paragraph("Accords : " + "  ".join(chords), S["meta"]))

    story.append(HRFlowable(width="100%", thickness=0.5, color=c_hr, spaceAfter=4*mm))

    sections_by_id = {s["id"]: s for s in song.get("sections", [])}
    sequence = [
        x for x in song.get("structure_sequence", [])
        if isinstance(x, str) and not x.startswith("_comment")
    ]
    sections_list = (
        [sections_by_id[sid] for sid in sequence if sid in sections_by_id]
        if sequence else song.get("sections", [])
    )

    for section in sections_list:
        label     = section.get("label", section.get("type", "Section"))
        chord_grid = section.get("chord_grid")
        lines     = section.get("lines", [])
        repeats   = section.get("repeats", 1)
        is_instr  = section.get("is_instrumental", False)

        elems = [Paragraph(f"[ {label.upper()} ]", S["section"])]

        if is_instr:
            perf_explicit = section.get("performance_progression")
            if perf_explicit and any(e.get("chords") for e in perf_explicit):
                # performance_progression prioritaire : groupements avec ×N par ligne
                for entry in perf_explicit:
                    chords = str(entry.get("chords") or "")
                    if not chords:
                        continue
                    suffix = f"  ×{entry['repeat']}" if entry.get("repeat", 1) > 1 else ""
                    elems.append(Paragraph(esc(chords + suffix), S["grid"]))
            elif chord_grid:
                for grid_line in chord_grid.split("\n"):
                    grid_line = grid_line.strip()
                    if grid_line:
                        suffix = f"   ×{repeats}" if repeats > 1 else ""
                        elems.append(Paragraph(esc(grid_line + suffix), S["grid"]))
            else:
                perf = _extract_performance_lines(section)
                has_chords = any(pl.get("chords") for pl in perf)
                if has_chords:
                    for pl in perf:
                        if not pl.get("chords"):
                            continue
                        suffix = f"  ×{pl['repeat']}" if pl.get("repeat", 1) > 1 else ""
                        elems.append(Paragraph(esc(pl["chords"] + suffix), S["grid"]))
                else:
                    elems.append(Paragraph("(instrumental)", S["instr"]))
        else:
            for line in lines:
                line_chords = line.get("chords", [])
                lyrics      = line.get("lyrics", "")
                if line_chords:
                    # Preformatted préserve les espaces de positionnement
                    elems.append(Preformatted(build_chord_line(line_chords, lyrics), S["chord"]))
                elems.append(Preformatted(lyrics or " ", S["lyric"]))
            if repeats > 1:
                elems.append(Paragraph(f"(× {repeats})", S["repeat"]))

        story.append(KeepTogether(elems))

    warnings = [w for w in song.get("warnings", []) if w.get("severity") in ("medium", "high")]
    if warnings:
        story.append(Spacer(1, 4*mm))
        story.append(Paragraph("Notes :", ParagraphStyle("nh", fontName="Helvetica-Bold", fontSize=9)))
        for w in warnings:
            story.append(Paragraph(f"• {esc(w['message'])}", S["warn"]))

    if "chords" not in _parts:
        story.clear()
    if "memo" in _parts:
        _add_memo_page_pdf(story, song, simplified_song=simplified_song, page_break=bool(story))

    doc = SimpleDocTemplate(
        str(output_path), pagesize=A4,
        topMargin=2*cm, bottomMargin=2*cm, leftMargin=2.5*cm, rightMargin=2*cm,
        title=f"{meta.get('title', '')} — {meta.get('artist', '')}",
    )
    doc.build(story)
    print(f"PDF généré  : {output_path}")


# ---------------------------------------------------------------------------
# Générateurs partiels (une page DOCX autonome)
# ---------------------------------------------------------------------------

def generate_chords_only(song: dict, output_path: Path):
    """Page 1 uniquement — fiche paroles + accords."""
    render_song, _ = _resolve_simplification(song)
    doc = Document()
    set_margins(doc, top=2.0, bottom=2.0, left=2.5, right=2.0)
    add_title_block(doc, render_song)
    sections_by_id = {s["id"]: s for s in render_song.get("sections", [])}
    sequence = [x for x in render_song.get("structure_sequence", [])
                if isinstance(x, str) and not x.startswith("_comment")]
    if sequence:
        for sid in sequence:
            sec = sections_by_id.get(sid)
            if sec:
                render_section(doc, sec)
    else:
        for sec in render_song.get("sections", []):
            render_section(doc, sec)
    doc.save(output_path)


def generate_memo_only(song: dict, output_path: Path):
    """Page 2 uniquement — fiche mémo structure guitare."""
    render_song, simplified_song = _resolve_simplification(song)
    doc = Document()
    set_margins(doc, top=2.0, bottom=2.0, left=2.5, right=2.0)
    add_memo_page_docx(doc, render_song, simplified_song=simplified_song, page_break=False)
    doc.save(output_path)


# ---------------------------------------------------------------------------
# Conversion DOCX → PDF via LibreOffice headless
# ---------------------------------------------------------------------------

def convert_docx_to_pdf(docx_path: Path, output_dir: Path) -> Path:
    """Convertit un DOCX en PDF via LibreOffice headless.
    Lève RuntimeError si LibreOffice est introuvable."""
    import subprocess
    candidates = [
        Path(r"C:\Program Files\LibreOffice\program\soffice.exe"),
        Path(r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"),
        Path("/usr/bin/soffice"),
        Path("/usr/lib/libreoffice/program/soffice"),
    ]
    soffice = next((p for p in candidates if p.exists()), None)
    if soffice is None:
        raise RuntimeError("LibreOffice introuvable — conversion DOCX→PDF impossible")
    output_dir.mkdir(parents=True, exist_ok=True)
    result = subprocess.run(
        [str(soffice), "--headless", "--convert-to", "pdf",
         "--outdir", str(output_dir), str(docx_path)],
        capture_output=True, text=True,
    )
    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice erreur : {result.stderr.strip()}")
    return output_dir / (docx_path.stem + ".pdf")


# ---------------------------------------------------------------------------
# Export PDF split (3 fichiers séparés)
# ---------------------------------------------------------------------------

_PART_LABELS = {
    "chords": "Paroles & Accords",
    "memo":   "Mémo Guitare",
}


def _split_pdf_filename(song: dict, part: str) -> str:
    """Retourne le nom de fichier PDF : 'Artiste - Titre - Type de fiche.pdf'."""
    meta = song["meta"]
    artist = meta.get("artist", "")
    title  = meta.get("title", "")
    label  = _PART_LABELS.get(part, part)

    def sanitize(s: str) -> str:
        for ch in r'\/:*?"<>|':
            s = s.replace(ch, "-")
        return s.strip()

    return f"{sanitize(artist)} - {sanitize(title)} - {label}.pdf"


def _generate_split_pdf_reportlab(render_song: dict, slug: str, simplified_song: dict = None):
    """Génère les 2 PDFs séparés via reportlab (fallback LibreOffice indisponible)."""
    PDF_EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    for part in ("chords", "memo"):
        generate_pdf(render_song, PDF_EXPORT_DIR / _split_pdf_filename(render_song, part),
                     simplified_song=simplified_song, parts=[part])


def _cleanup_deprecated_pdfs(song: dict) -> None:
    """Supprime les PDFs de fiches retirées du workflow si présents sur disque."""
    meta = song.get("meta", {})
    artist = meta.get("artist", "")
    title  = meta.get("title", "")

    def sanitize(s: str) -> str:
        for ch in r'\/:*?"<>|':
            s = s.replace(ch, "-")
        return s.strip()

    for old_label in ("Comprendre le Morceau",):
        p = PDF_EXPORT_DIR / f"{sanitize(artist)} - {sanitize(title)} - {old_label}.pdf"
        if p.exists():
            p.unlink()


def generate_split_pdf(song: dict, slug: str):
    """Orchestrateur : génère 2 DOCX partiels, les convertit en PDF via LibreOffice,
    ou tombe sur reportlab si LibreOffice est indisponible."""
    import shutil
    render_song, simplified_song = _resolve_simplification(song)
    PDF_EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    _cleanup_deprecated_pdfs(song)

    tmp = OUTPUT_DIR / "_split_tmp"
    tmp.mkdir(exist_ok=True)

    generators = {
        "chords": generate_chords_only,
        "memo":   generate_memo_only,
    }

    try:
        for part, gen in generators.items():
            stem = _split_pdf_filename(song, part)[:-4]  # retire .pdf
            docx = tmp / f"{stem}.docx"
            gen(song, docx)
            convert_docx_to_pdf(docx, PDF_EXPORT_DIR)
        shutil.rmtree(tmp, ignore_errors=True)
    except RuntimeError as e:
        shutil.rmtree(tmp, ignore_errors=True)
        print(f"Info : {e} — utilisation de reportlab.")
        _generate_split_pdf_reportlab(render_song, slug, simplified_song=simplified_song)

    print(f"PDFs exportés dans : {PDF_EXPORT_DIR}")


# ---------------------------------------------------------------------------
# Point d'entrée
# ---------------------------------------------------------------------------

def main():
    import argparse
    _ensure_utf8()

    parser = argparse.ArgumentParser(
        description="Génère le DOCX/PDF d'une fiche de chords.")
    parser.add_argument("json_path", help="Chemin vers le JSON song validé")
    parser.add_argument("--split-pdf", action="store_true",
                        help="Exporte 2 PDFs séparés dans PDF_EXPORT_DIR")
    parser.add_argument("--part", choices=["chords", "memo"],
                        help="Exporte un seul PDF dans PDF_EXPORT_DIR")
    args = parser.parse_args()

    song = load_song(args.json_path)
    check_validated(song)

    slug = song["meta"].get("slug", "output")
    OUTPUT_DIR.mkdir(exist_ok=True)
    output_path = OUTPUT_DIR / f"song_{slug}.docx"

    render_song, simplified_song = _resolve_simplification(song)
    generate(song, output_path)

    # PDF complet dans OUTPUT_DIR
    try:
        pdf_path = OUTPUT_DIR / f"song_{slug}.pdf"
        generate_pdf(render_song, pdf_path, simplified_song=simplified_song)
    except Exception as e:
        print(f"Avertissement PDF : {e}")

    # Export split
    if args.split_pdf:
        generate_split_pdf(song, slug)
    elif args.part:
        PDF_EXPORT_DIR.mkdir(parents=True, exist_ok=True)
        pdf_path = PDF_EXPORT_DIR / _split_pdf_filename(song, args.part)
        try:
            generate_pdf(render_song, pdf_path, simplified_song=simplified_song,
                         parts=[args.part])
        except Exception as e:
            print(f"Avertissement PDF : {e}")


if __name__ == "__main__":
    main()
